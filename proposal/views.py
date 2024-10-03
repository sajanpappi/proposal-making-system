# Ensure this function is present in views.py
from django.shortcuts import render, redirect
from .utils import generate_proposal

def index(request):
    if request.method == 'POST':
        # Collect form inputs from Screen 1
        project_summary = request.POST.get('project_summary')
        client_name = request.POST.get('client_name')

        # Store data in session
        request.session['project_summary'] = project_summary
        request.session['client_name'] = client_name

        # Redirect to the branding and styling screen
        return redirect('branding_styling')

    return render(request, 'proposal/index.html')


def branding_styling(request):
    if request.method == 'POST':
        # Collect form data from Screen 2
        has_logo = request.POST.get('has_logo')
        primary_color = request.POST.get('primary_color')
        secondary_color = request.POST.get('secondary_color')
        website_style = request.POST.get('website_style')

        # Store in session
        request.session['has_logo'] = has_logo
        request.session['primary_color'] = primary_color
        request.session['secondary_color'] = secondary_color
        request.session['website_style'] = website_style

        # Redirect to the next screen
        return redirect('pages_features')

    return render(request, 'proposal/branding_styling.html')

def pages_features(request):
    if request.method == 'POST':
        # Collect form data from Screen 3
        website_goal = request.POST.get('website_goal')
        visitor_action = request.POST.get('visitor_action')
        target_audience = request.POST.get('target_audience')
        benchmark_website = request.POST.get('benchmark_website')

        # Store in session
        request.session['website_goal'] = website_goal
        request.session['visitor_action'] = visitor_action
        request.session['target_audience'] = target_audience
        request.session['benchmark_website'] = benchmark_website

        # Redirect to Screen 4
        return redirect('associated_services')

    return render(request, 'proposal/pages_features.html')

def associated_services(request):
    if request.method == 'POST':
        # Collect form data from Screen 4
        need_domain = request.POST.get('need_domain')
        need_hosting = request.POST.get('need_hosting')
        need_seo = request.POST.get('need_seo')

        # Store in session
        request.session['need_domain'] = need_domain
        request.session['need_hosting'] = need_hosting
        request.session['need_seo'] = need_seo

        # Redirect to Screen 5 (Budget)
        return redirect('budget')

    return render(request, 'proposal/associated_services.html')

def budget(request):
    if request.method == 'POST':
        # Collect form data from Screen 5
        budget_range = request.POST.get('budget_range')

        # Store in session
        request.session['budget_range'] = budget_range

        # Redirect to the proposal generation
        return redirect('generate_proposal')

    return render(request, 'proposal/budget.html')

def generate_proposal_view(request):
    # Retrieve session data
    project_summary = request.session.get('project_summary', '')
    client_name = request.session.get('client_name', '')
    has_logo = request.session.get('has_logo', '')
    primary_color = request.session.get('primary_color', '')
    secondary_color = request.session.get('secondary_color', '')
    website_style = request.session.get('website_style', '')
    website_goal = request.session.get('website_goal', '')
    visitor_action = request.session.get('visitor_action', '')
    target_audience = request.session.get('target_audience', '')
    benchmark_website = request.session.get('benchmark_website', '')
    need_domain = request.session.get('need_domain', '')
    need_hosting = request.session.get('need_hosting', '')
    need_seo = request.session.get('need_seo', '')
    budget_range = request.session.get('budget_range', '')

    # Build prompt based on collected data
    prompt = f"""
    About CreaveLabs:
    CreaveLabs is a web development company which delivers you a web experience with international standards. Let us be your technical partner, leave the rest to us. We analyze your business, elicit your requirements and develop the websites which exactly suit your purpose. Our technical, creative and managerial expertise can ensure you a web experience rather than a website.

    Project Description:
    {project_summary}

    Branding and Styling:
    Logo Preference: {has_logo}
    Primary Colors: {primary_color}
    Secondary Colors: {secondary_color}
    Website Style: {website_style}

    Pages and Features:
    Primary Goal: {website_goal}
    Expected Action: {visitor_action}
    Target Audience: {target_audience}
    Benchmark Website: {benchmark_website}

    Associated Services:
    Domain: {need_domain}
    Hosting: {need_hosting}
    SEO: {need_seo}

    Budget Range:
    {budget_range}

    You act as the Business Analyst of a web design agency named CreaveLabs and Please generate a detailed proposal based on the above information. In the generated 
    proposal, please include the following details: Contact Person from CreaveLabs will be Mr. Sajan Pappachen, He is the business Director of the company and his 
    phone number is +919995997334. Also, in the generated proposal for all Main titles like Introduction, Project understanding, Proposed Solution, Wesbite Design and Development, 
    Design, Pages, Functionality, Technical Specifications, Project Timelines, Budget and Payment Terms, CreaveLabs Advantage, Next Steps etc use Astericks. For 
    example, "Introduction:". For sub titles that fall under these main titles , like page names, functionality names, platform, hosting, domain, SEO, Phases under timeline like Phase1, Estimated 
    budget, Payment Terms that falls under the main title Budget and Payment Terms please add a hyphen after the title. For example "Home Page-",  "Hosting- ", "Domain-". Also, note that the 
    platfomrm that CreaveLabs uses for website development will be Laravel and Bootstrap. In budget calculation, if the client want a domain, give the pricing of domain as INR 1150 and if the client need hosting, give its pricing as 
    INR 3500 for 1 GB space. If the client want SEO, please give the pricing as INR 35000 per month for 4  Keywords. For website design and development pricing, give 
    an average between maximum and minimum budget provided by the client. The pricing of Domain, Hosting and SEO will not be included in the pricing of website design and development. 
    """

    try:
        # Generate proposal
        proposal_text = generate_proposal(prompt)
        
        # Handle empty or error responses
        if not proposal_text:
            raise ValueError("Generated proposal is empty or failed.")

    except Exception as e:
        # Log the exception or handle it as needed
        print(f"Error generating proposal: {e}")
        return render(request, 'proposal/error.html', {'error_message': 'An error occurred while generating the proposal. Please try again later.'})

    # Render proposal display page
    return render(request, 'proposal/display_proposal.html', {
        'proposal_text': proposal_text
    })

from django.http import HttpResponse
from docx import Document
from io import BytesIO

from django.http import HttpResponse
from docx import Document
from docx.shared import Pt
from io import BytesIO

from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from docx.shared import Inches
import os

def download_proposal(request):
    # Retrieve the proposal text from the POST request
    proposal_text = request.POST.get('proposal_text', '')

    # Check if proposal text is present
    if not proposal_text:
        return HttpResponse("No proposal text provided.", status=400)

    # Create a Word document
    doc = Document()

    # 1. Add an image at the beginning (full width)
    # Adjust the image path according to your project structure
    image_path = os.path.join('path_to_images_folder', 'header_image.jpg')  # Replace with your image path
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(6))  # Full-width image

    # Add and style the title: "CreaveLabs Proposal for [Client Name]"
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(f"CreaveLabs Proposal for {request.session.get('client_name', 'Client Name')} Website")
    title_run.font.size = Pt(20)  # Set font size
    title_run.font.bold = True  # Make font bold
    title_run.font.color.rgb = RGBColor(0, 112, 192)  # Blue color
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align the title

    # Add Date section: "Date: [Today's Date]"
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(f"Date: {request.session.get('date', "Today\'s Date")}")
    date_run.font.size = Pt(12)

    # Add Client section: "Client: [Client Name]"
    client_paragraph = doc.add_paragraph()
    client_run = client_paragraph.add_run(f"Client: {request.session.get('client_name', 'Client Name')}")
    client_run.font.size = Pt(12)
    client_run.font.bold = True
    client_run.font.color.rgb = RGBColor(0, 112, 192)  # Blue color

    # Add Project section: "Project: [Project Name]"
    project_paragraph = doc.add_paragraph()
    project_run = project_paragraph.add_run(f"Project: {request.session.get('project_name', 'Project Name')}")
    project_run.font.size = Pt(12)
    project_run.font.bold = True
    project_run.font.color.rgb = RGBColor(0, 112, 192)  # Blue color

    # 2. Remove asterisks from the proposal text
    proposal_text = proposal_text.replace('*', '')

    # Add the generated proposal text to the Word document
    paragraphs = proposal_text.split('<br>')  # Split based on <br> (or use your own delimiter)
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph.strip())

    # 3. Add company details at the end
    company_details = doc.add_paragraph("Sincerely,\nCreaveLabs Team")
    company_website = doc.add_paragraph("Website: www.creavelabs.com")
    company_email = doc.add_paragraph("Email: mail@creavelabs.com")

    # 4. Add CreaveLabs logo at the end of the document
    logo_path = os.path.join('path_to_images_folder', 'creavelabs_logo.png')  # Replace with the actual logo path
    if os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Inches(2))  # Adjust size as needed
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align the logo

    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Create HTTP response
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=proposal_{request.session.get("client_name", "proposal")}.docx'
    return response

